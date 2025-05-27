import os
from typing import List, Dict
import google.generativeai as genai
from langchain_google_genai import GoogleGenerativeAIEmbeddings, ChatGoogleGenerativeAI
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain.chains import ConversationalRetrievalChain
from langchain.memory import ConversationBufferMemory
from langchain.schema import Document
import PyPDF2
try:
    from docx import Document as DocxDocument
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False
from config import Config

# Configure Gemini API
GOOGLE_API_KEY = Config.GEMINI_API_KEY
genai.configure(api_key=GOOGLE_API_KEY)

class FoodReviewRAG:
    def __init__(self, documents_dir: str):
        self.documents_dir = documents_dir
        self.embeddings = GoogleGenerativeAIEmbeddings(
            model="models/embedding-001",
            google_api_key=GOOGLE_API_KEY
        )
        self.llm = ChatGoogleGenerativeAI(
            model="models/gemini-1.5-flash",
            temperature=0.7,
            google_api_key=GOOGLE_API_KEY
        )
        self.vector_store = None
        self.qa_chain = None
        self.memory = ConversationBufferMemory(
            memory_key="chat_history",
            return_messages=True,
            output_key="answer"
        )

    def load_documents(self) -> List[Document]:
        """Load documents from the specified directory."""
        documents = []
        for filename in os.listdir(self.documents_dir):
            file_path = os.path.join(self.documents_dir, filename)
            if filename.endswith('.txt'):
                with open(file_path, 'r', encoding='utf-8') as file:
                    text = file.read()
                    documents.append(Document(page_content=text, metadata={"source": filename}))
            elif filename.endswith('.pdf'):
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for i, page in enumerate(pdf_reader.pages):
                        text = page.extract_text()
                        documents.append(Document(page_content=text, metadata={"source": f"{filename} - page {i+1}"}))
            elif filename.endswith('.docx') and DOCX_SUPPORT:
                doc = DocxDocument(file_path)
                text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                documents.append(Document(page_content=text, metadata={"source": filename}))
            elif filename.endswith('.docx') and not DOCX_SUPPORT:
                print(f"Warning: .docx support is not available. Skipping {filename}")
        return documents

    def process_documents(self):
        """Process documents and create vector store."""
        # Load documents
        documents = self.load_documents()
        
        # Split documents into chunks
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200
        )
        chunks = text_splitter.split_documents(documents)
        
        # Create vector store
        self.vector_store = FAISS.from_documents(chunks, self.embeddings)
        
        # Create QA chain
        self.qa_chain = ConversationalRetrievalChain.from_llm(
            llm=self.llm,
            retriever=self.vector_store.as_retriever(),
            memory=self.memory,
            return_source_documents=True
        )

    def query(self, question: str) -> Dict:
        """Query the RAG system with a question."""
        if not self.qa_chain:
            raise ValueError("Please process documents first by calling process_documents()")
        
        result = self.qa_chain({"question": question})
        return {
            "answer": result["answer"],
            "source_documents": result["source_documents"]
        }

def main():
    # Initialize RAG system
    rag = FoodReviewRAG("foodreview")
    
    # Process documents
    print("Processing documents...")
    rag.process_documents()
    print("Documents processed successfully!")
    
    # Interactive query loop
    print("\nAsk questions about food reviews (type 'quit' to exit):")
    while True:
        question = input("\nYour question: ")
        if question.lower() == 'quit':
            break
            
        try:
            result = rag.query(question)
            print("\nAnswer:", result["answer"])
            print("\nSources:")
            for doc in result["source_documents"]:
                print("-", doc.page_content[:200], "...")
        except Exception as e:
            print(f"Error: {str(e)}")

if __name__ == "__main__":
    main() 